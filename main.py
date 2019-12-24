#!/usr/bin/env python3

from os import system, name
import sys
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
import pandas as pd


# Functions
def clear_screen():
    if name == 'nt':
        _ = system('cls')
    else:
        _ = system('clear')


def head():
    print('+-+ +-+ +-+ +-+ +-+   +-+ +-+ +-+ +-+ +-+ +-+')
    print('|B| |L| |A| |C| |K|   |C| |o| |d| |i| |n| |g|')
    print('+-+ +-+ +-+ +-+ +-+   +-+ +-+ +-+ +-+ +-+ +-+')
    print('\n' + '-' * 45)
    print('Program: Docx.Table Analysis')


def homemenu():
    clear_screen()
    head()
    print('\n' + '-' * 45)
    print('Commands:')
    print('[1] Read Docx.Table')
    # print('[2] test')
    # print('[3] test')
    # print('[4] test')
    # print('[5] test')
    # print('\n')
    print('[9] Quit program')
    print('-' * 45)
    input_ = input('>>> ')

    # Quit program
    if input_ == '9':
        print('Program terminated.')
        sys.exit(1)

    # Read Docx.Table
    elif input_ == '1':
        submenu_1()

    # Break
    else:
        print('No valid choice. Program terminated.')
        sys.exit(1)


def submenu_1():
    while True:
        info_ = 'no further info'
        clear_screen()
        head()
        print('Submenu: Read Docx.Table')
        print('-' * 45)
        print(info_)
        print('-' * 45)
        file_ = input('Enter file-name: >>> ')

        # doc = Document(file_)

        try:
            doc = Document(file_)
        except PackageNotFoundError:
            print('File not found. Program terminated.')
            sys.exit(1)

        while True:
            tables = doc.tables
            info_ = ('File: ' + file_ + '\n' + 'Number of tables: ' + str(len(tables)))
            clear_screen()
            head()
            print('Submenu: Read Docx.Table')
            print('-' * 45)
            print(info_)
            print('-' * 45)
            table_ = int(input('Enter table-number: >>> ')) - 1

            result = []

            try:
                for row in tables[table_].rows:
                    interim = []
                    result.append(interim)
                    for cell in row.cells:
                        interim.append(cell.text)
            except IndexError:
                print('Table not found. Program terminated.')
                sys.exit(1)

            info_ = (
                    'File: ' + file_ + '\n'
                    + 'Number of tables: ' + str(len(tables)) + '\n'
                    + 'Selected table: ' + str((table_ + 1))
                    )
            clear_screen()
            head()
            print('Submenu: Read Docx.Table')
            print('-' * 45)
            print(info_)
            print('-' * 45)

            print('\n>>> Printing Dataframe <<<\n')
            # print(result)
            labels = result[0]
            df = pd.DataFrame.from_records(result[1:], columns=labels)

            print(df)
            print('\n')

            head()
            print('Submenu: Read Docx.Table')
            print('-' * 45)
            print(info_)
            print('-' * 45)
            loop_ = input('Type "home" for home-screen or "exit" for terminating the program: >>> ')

            if loop_ == 'home':
                homemenu()

            elif loop_ == 'exit':
                print('Program terminated.')
                sys.exit(1)

            else:
                print('No valid choice. Program terminated.')
                sys.exit(1)


# Main
while True:
    homemenu()
